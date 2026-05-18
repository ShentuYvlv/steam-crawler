from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("/Users/zed/all code/D 互影/steam-crawler/软著")
TEMPLATE = ROOT / "情感反诈模拟器-操作手册:说明书.docx"
OUTPUT = ROOT / "线上社区运营管理系统-用户操作手册.docx"
SCREENSHOT_DIR = ROOT / "screenshots" / "framed"

SOFTWARE_NAME = "线上社区运营管理系统"
VERSION = "V1.0"
MANUAL_NAME = f"{SOFTWARE_NAME} {VERSION} 用户操作手册"
BODY_FONT = "微软雅黑"


def clear_document(doc: Document) -> None:
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def set_run_font(run, name: str, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)

    for key in ["w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"]:
        r_fonts.set(qn(key), name)

    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")
    lang.set(qn("w:bidi"), "ar-SA")


def set_style_font(style, name: str, size: float, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold

    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)

    for key in ["w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"]:
        r_fonts.set(qn(key), name)

    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")
    lang.set(qn("w:bidi"), "ar-SA")


def reset_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    for container in [section.header, section.footer]:
        element = container._element
        for child in list(element):
            element.remove(child)

    header_para = section.header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(MANUAL_NAME)
    set_run_font(run, BODY_FONT, 10.5, bold=False)

    footer_para = section.footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(MANUAL_NAME)
    set_run_font(run, BODY_FONT, 10.5, bold=False)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, BODY_FONT, 11, bold=False)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.5

    heading1 = doc.styles["Heading 1"]
    set_style_font(heading1, BODY_FONT, 18, bold=True)
    heading1.paragraph_format.space_before = Pt(24)
    heading1.paragraph_format.space_after = Pt(0)

    heading2 = doc.styles["Heading 2"]
    set_style_font(heading2, BODY_FONT, 15, bold=True)
    heading2.paragraph_format.space_before = Pt(10)
    heading2.paragraph_format.space_after = Pt(0)

    list_bullet = doc.styles["List Bullet"]
    set_style_font(list_bullet, BODY_FONT, 11, bold=False)


def add_center_text(doc: Document, text: str, size: float, bold: bool = False) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_run_font(run, BODY_FONT, size, bold=bold)


def add_blank(doc: Document, count: int = 1) -> None:
    for _ in range(count):
        doc.add_paragraph("")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_paragraph(style=f"Heading {level}")
    run = para.add_run(text)
    set_run_font(run, BODY_FONT, 18 if level == 1 else 15, bold=True)


def add_paragraph(doc: Document, text: str, center: bool = False) -> None:
    para = doc.add_paragraph(style="Normal")
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_run_font(run, BODY_FONT, 11, bold=False)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        set_run_font(run, BODY_FONT, 11, bold=False)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_cover(doc: Document) -> None:
    add_blank(doc, 6)
    add_center_text(doc, SOFTWARE_NAME, 36, bold=True)
    add_blank(doc, 1)
    add_center_text(doc, "用  户  操  作  手  册", 22)
    add_blank(doc, 3)
    add_center_text(doc, f"版本号：{VERSION}", 14)
    add_blank(doc, 1)
    add_center_text(doc, "2026年5月", 14)


def add_toc_page(doc: Document, title: str, lines: list[str]) -> None:
    add_heading(doc, title, 1)
    add_blank(doc, 1)
    for line in lines:
        add_paragraph(doc, line)


def add_image_page(doc: Document, chapter: str, section: str, intro: str, image_name: str, caption: str) -> None:
    add_heading(doc, chapter, 1)
    add_heading(doc, section, 2)
    add_paragraph(doc, intro)
    add_blank(doc, 1)
    doc.add_picture(str(SCREENSHOT_DIR / image_name), width=Cm(15.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, f"图：{caption}", center=True)
    add_paragraph(doc, "上图为系统当前模块的标准操作界面截图，截图内容完整保留了浏览器窗口头与软件名称。", center=False)


def add_text_page(doc: Document, section: str, summary: str, bullets: list[str], note: str | None = None) -> None:
    add_heading(doc, section, 2)
    add_paragraph(doc, summary)
    add_bullets(doc, bullets)
    if note:
        add_paragraph(doc, note)


def build_pages(doc: Document) -> None:
    add_cover(doc)
    add_page_break(doc)

    toc_page_1 = [
        "1  软件概述...............................................4",
        "2  登录与身份验证.........................................5",
        "3  运营统计总览...........................................7",
        "4  游戏列表...............................................9",
        "5  游戏编辑与分类........................................11",
        "6  评论列表..............................................13",
        "7  回复策略管理..........................................15",
        "8  任务与同步............................................17",
        "9  任务队列..............................................19",
    ]
    add_toc_page(doc, "目  录", toc_page_1)
    add_page_break(doc)

    toc_page_2 = [
        "10  回复记录.............................................21",
        "11  用户管理.............................................23",
        "12  数据筛选与导出.......................................25",
        "13  自有产品与竞品管理规则...............................26",
        "14  权限与账号安全.......................................27",
        "15  异常处理与日常维护...................................28",
        "16  操作规范与最佳实践...................................29",
        "17  结语.................................................30",
    ]
    add_toc_page(doc, "目  录（续）", toc_page_2)
    add_page_break(doc)

    add_heading(doc, "第1章  软件概述", 1)
    add_heading(doc, "1.1  产品简介", 2)
    add_paragraph(doc, f"{SOFTWARE_NAME}是一套面向评论处理、内容审核、回复协作、任务同步与数据统计场景的运营后台系统。系统通过统一的工作台将游戏清单、评论列表、回复策略、审核记录、任务队列与用户权限集中到同一平台内，帮助运营人员在单一界面中完成日常管理工作。")
    add_paragraph(doc, "系统采用前后端分离方式构建，支持多游戏管理、自有产品与竞品分流管理、评论筛选与导出、任务同步跟踪、回复草稿审核与发送记录留存等能力，适合日常评论运营、客服协同与内容复盘场景。")
    add_heading(doc, "1.2  主要能力", 2)
    add_bullets(doc, [
        "统一管理自有产品与竞品条目，并按分类控制可用能力。",
        "提供评论总览、趋势分析、状态分布和回复结果等统计视图。",
        "支持评论筛选、详情查看、状态标记、导出与审核处理。",
        "支持回复策略版本管理、任务同步配置、队列日志和用户权限管理。",
    ])
    add_paragraph(doc, "操作手册以下各章将按照实际界面顺序依次介绍主要模块的进入方式、界面组成和常见操作。")
    add_page_break(doc)

    add_image_page(
        doc,
        "第2章  登录与身份验证",
        "2.1  登录界面",
        "启动系统后，浏览器会首先进入登录页面。操作人员在此输入账号和密码，通过系统的统一认证后进入后台工作区。",
        "01-login-page.png",
        "登录界面",
    )
    add_page_break(doc)

    add_text_page(
        doc,
        "2.2  操作说明",
        "登录模块用于确认用户身份并加载对应权限范围，成功登录后系统会自动保持会话状态，后续页面切换无需重复输入凭据。",
        [
            "在“用户名”输入框中填写管理员或运营账号。",
            "在“密码”输入框中输入对应密码后点击“登录”按钮。",
            "若输入内容不完整，按钮保持禁用状态，避免误提交。",
            "登录成功后会自动进入运营总览页，并加载当前用户的角色信息。",
            "如需结束当前会话，可在左下角账号卡片中点击“退出登录”。",
        ],
        "建议由管理员统一创建和维护账号，不开放匿名注册，以保证系统内数据与操作记录的一致性。",
    )
    add_page_break(doc)

    add_image_page(
        doc,
        "第3章  运营统计总览",
        "3.1  总览首页",
        "登录成功后默认进入运营统计总览页。该页面展示评论总数、好评率、已回复数量、回复成功率以及近 14 天趋势图，是进入日常工作前的总览面板。",
        "02-dashboard.png",
        "运营统计总览页",
    )
    add_page_break(doc)

    add_text_page(
        doc,
        "3.2  界面说明",
        "总览页用于帮助运营人员快速判断当前处理压力和回复效果，同时支持按自有游戏维度切换统计范围。",
        [
            "左侧导航栏显示全部核心模块，点击可在各页面间切换。",
            "顶部总览范围下拉框可切换“全部自家游戏”或指定游戏。",
            "统计卡片用于展示评论总量、当前好评率、已回复数和回复成功率。",
            "处理状态面板展示待处理、已忽略、好评和差评等数量分布。",
            "趋势图区域展示近 14 天评论新增量与回复发送量，适合日常观察波动。",
        ],
        "建议每日开始工作时先查看总览页，以确认需要优先处理的模块和当天的工作节奏。",
    )
    add_page_break(doc)

    add_image_page(
        doc,
        "第4章  游戏列表",
        "4.1  游戏清单",
        "游戏列表页用于维护运营对象，支持查看游戏名称、分类、评论数、监控状态、最近同步结果以及手动同步入口。",
        "03-games-list.png",
        "游戏列表页",
    )
    add_page_break(doc)

    add_text_page(
        doc,
        "4.2  操作说明",
        "游戏列表页同时承担“查看清单”和“发起同步”的作用，适合作为运营对象总表使用。",
        [
            "点击“新增游戏”可录入新的监控条目。",
            "点击“全部同步”可一次性提交所有游戏的同步任务。",
            "列表中的“自家/竞品”标签用于标识不同的管理范围。",
            "点击单行右侧“同步”按钮可单独触发某个游戏的手动同步。",
            "点击“编辑”按钮可在右侧表单中修改该游戏的名称、分类和同步配置。",
        ],
        "如果需要将某个游戏纳入回复运营流程，应在编辑表单中将其分类设置为“自家游戏”。",
    )
    add_page_break(doc)

    add_image_page(
        doc,
        "第5章  游戏编辑与分类",
        "5.1  编辑表单",
        "在游戏列表页向下滚动后，可以看到游戏编辑区域。该区域支持录入游戏名称、分类、同步策略和每日执行时间等信息。",
        "04-games-edit.png",
        "游戏编辑与分类表单",
    )
    add_page_break(doc)

    add_text_page(
        doc,
        "5.2  表单说明",
        "编辑表单用于维护单个游戏的基础配置，完成保存后会立即影响后续统计范围和任务行为。",
        [
            "“游戏分类”用于区分自有产品与竞品，竞品默认只开放查看和导出能力。",
            "“同时维护这款游戏的监控任务”用于控制是否一并配置同步规则。",
            "执行时间、语言、过滤方式、购买类型等字段决定后续同步策略。",
            "保存后系统会同步刷新游戏清单和相关任务配置。",
            "若仅需保留游戏档案信息，可以关闭定时同步并只维护基础字段。",
        ],
        "建议对自有产品配置稳定的每日同步时间，并保持分类准确，避免运营能力误用到竞品。",
    )
    add_page_break(doc)

    add_image_page(
        doc,
        "第6章  评论列表",
        "6.1  评论主列表",
        "评论列表页是日常处理评论的核心页面。系统支持按游戏范围、评论分组、关键词、状态、时长和互动指标进行多条件筛选。",
        "05-reviews-list.png",
        "评论列表页",
    )
    add_page_break(doc)

    add_text_page(
        doc,
        "6.2  使用说明",
        "评论列表页面向运营与审核角色提供统一处理视图，支持在同一页面完成筛选、浏览、导出与状态管理。",
        [
            "可切换“自家游戏”和“竞品”两个视图，默认进入自家游戏视图。",
            "可切换“待处理”和“已回复”两个分组，用于聚焦当前工作内容。",
            "顶部筛选区支持关键词、状态、回复状态、时间范围和时长范围组合查询。",
            "支持“导出当前”和“导出全部”，便于离线分析与归档。",
            "竞品视图默认只开放查看、筛选和导出，不提供回复相关写操作。",
        ],
        "建议在处理大量评论时先缩小游戏范围，再叠加分组和关键词，以提高筛选效率。",
    )
    add_page_break(doc)

    add_image_page(
        doc,
        "第7章  回复策略管理",
        "7.1  Skill 配置页",
        "回复策略管理页面用于维护评论回复所依赖的策略文档、模型名称、温度参数与版本状态，适合运营与内容负责人协同调整策略。",
        "06-reply-strategies.png",
        "回复策略管理页",
    )
    add_page_break(doc)

    add_text_page(
        doc,
        "7.2  操作说明",
        "策略页通过版本化方式管理回复能力，可在不影响历史记录的前提下持续演进策略。",
        [
            "左侧列表用于查看历史 Skill 版本和当前激活版本。",
            "点击右上角新增按钮可创建新的回复策略文档。",
            "可维护 Skill 名称、模型、温度、描述和完整策略正文。",
            "未激活的版本可通过“设为 Active”切换为当前生效版本。",
            "保存后系统会将新版本用于后续草稿生成，并保留历史版本快照。",
        ],
        "建议每次调整策略时记录修改目的，例如语气优化、风险词规避或新活动话术接入，以便后续复盘。",
    )
    add_page_break(doc)

    add_image_page(
        doc,
        "第8章  任务与同步",
        "8.1  监控任务页面",
        "任务与同步页面用于配置各游戏的定时同步任务，支持新建、编辑、删除和手动同步操作。",
        "07-tasks.png",
        "任务与同步页",
    )
    add_page_break(doc)

    add_text_page(
        doc,
        "8.2  操作说明",
        "监控任务页面适合维护日常同步计划，并结合任务列表观察最近运行情况。",
        [
            "左侧卡片展示当前所有监控任务，可按任务快速切换。",
            "右侧表单可编辑任务名称、对应游戏、执行时间及同步参数。",
            "支持直接点击“手动同步”立即提交一次同步任务。",
            "停用任务后系统不再按计划执行，但保留配置内容。",
            "删除任务会移除该条计划配置，通常仅在监控对象废弃时使用。",
        ],
        "建议自有产品采用固定时段同步，避免频繁触发造成日志碎片化。",
    )
    add_page_break(doc)

    add_image_page(
        doc,
        "第9章  任务队列",
        "9.1  队列与详情",
        "任务队列页展示当前正在运行、排队中和已结束的任务，并支持查看详细日志、代理状态和执行结果。",
        "08-task-queue.png",
        "任务队列页",
    )
    add_page_break(doc)

    add_text_page(
        doc,
        "9.2  使用说明",
        "队列页主要用于排查同步任务运行状态，便于运营和技术人员定位异常。",
        [
            "可在左侧切换“进行中与排队”或“已结束任务”两类视图。",
            "支持按游戏筛选任务列表，快速定位指定游戏的执行记录。",
            "点击某条任务后，可查看执行数量、开始时间、结束时间和任务日志。",
            "若任务允许取消，系统会在详情区提供取消入口。",
            "代理状态卡片用于辅助判断抓取链与发送链的网络出口情况。",
        ],
        "出现同步失败时，建议优先查看任务详情中的错误信息与日志，再决定是否重新发起同步。",
    )
    add_page_break(doc)

    add_image_page(
        doc,
        "第10章  回复记录",
        "10.1  审核与发送记录",
        "回复记录页面按照游戏与处理日期聚合待审核草稿和已发送记录，用于日常审核、回溯和删除需求登记。",
        "09-reply-records.png",
        "回复记录页",
    )
    add_page_break(doc)

    add_text_page(
        doc,
        "10.2  使用说明",
        "该页面采用四列结构，便于从游戏、日期、草稿与已发记录四个角度查看回复处理过程。",
        [
            "左侧第一列为游戏切换区，可选择当前要查看的游戏。",
            "第二列按处理日期聚合记录，便于查看某一天的审核与发送结果。",
            "第三列展示待审核草稿，可进行保存、重生成、驳回与通过发送。",
            "第四列展示已发送回复，可查看原评论链接并登记删除需求。",
            "页面仅展示自有产品记录，竞品不会进入回复审核与发送链路。",
        ],
        "建议审核人员每日处理后回到本页面复核发送结果，及时标记失败原因或后续动作。",
    )
    add_page_break(doc)

    add_image_page(
        doc,
        "第11章  用户管理",
        "11.1  账号维护",
        "用户管理页面仅对管理员开放，用于创建新账号、调整角色以及启用或停用账户。",
        "10-users.png",
        "用户管理页",
    )
    add_page_break(doc)

    add_text_page(
        doc,
        "11.2  操作说明",
        "用户管理页用于保证系统权限边界清晰，适合管理员进行统一授权和账号维护。",
        [
            "顶部表单可直接创建新的后台账号。",
            "支持维护显示名、角色类型、密码和启用状态。",
            "列表区域可对现有用户进行角色切换和启停操作。",
            "管理员账号通常用于系统配置，运营账号用于日常处理，查看账号仅用于只读访问。",
            "账号创建后建议立即验证登录，并在离岗或角色变化时及时调整权限。",
        ],
        "密码建议由管理员统一下发和轮换，不建议多个岗位共用同一账号。",
    )
    add_page_break(doc)

    add_heading(doc, "第12章  数据筛选与导出", 1)
    add_heading(doc, "12.1  常用导出场景", 2)
    add_paragraph(doc, "系统提供“导出当前”和“导出全部”两种导出模式。前者基于当前筛选条件输出结果，适合用于专项复盘；后者基于当前游戏维度输出全量数据，适合做周报、月报或离线分析。")
    add_bullets(doc, [
        "在评论列表页完成筛选后，可直接导出当前结果集。",
        "切换至竞品视图后，同样可以导出竞品评论数据。",
        "导出文件适合用于表格统计、情绪分析和人工抽样复核。",
        "执行导出前建议先确认分组、游戏范围和关键词条件，避免误导出。",
    ])
    add_page_break(doc)

    add_heading(doc, "第13章  自有产品与竞品管理规则", 1)
    add_heading(doc, "13.1  分类原则", 2)
    add_paragraph(doc, "系统已将运营对象拆分为“自有产品”和“竞品”两类。自有产品可进入评论处理、草稿生成、回复审核与发送链路；竞品仅保留查看、筛选、详情和导出能力。")
    add_bullets(doc, [
        "新增游戏默认按竞品处理，避免误开放写操作。",
        "只有明确需要运营的对象才应切换为自有产品。",
        "回复记录、审核队列和发送结果仅展示自有产品数据。",
        "统计总览默认只汇总自有产品范围，避免竞品干扰日常判断。",
    ])
    add_page_break(doc)

    add_heading(doc, "第14章  权限与账号安全", 1)
    add_heading(doc, "14.1  权限建议", 2)
    add_paragraph(doc, "系统支持管理员、运营和只读等不同角色。建议按照岗位分配最小必要权限，并结合账号启停功能控制访问范围。")
    add_bullets(doc, [
        "管理员负责账号、策略、系统级配置和关键参数维护。",
        "运营人员负责评论处理、审核和日常同步操作。",
        "只读人员适合用于查看统计报表、评论详情与导出数据。",
        "涉及高风险操作时应使用单独账号并保留操作记录。",
    ])
    add_page_break(doc)

    add_heading(doc, "第15章  异常处理与日常维护", 1)
    add_heading(doc, "15.1  常见处理建议", 2)
    add_paragraph(doc, "当系统出现数据为空、任务失败、导出异常或登录异常时，可按照“先看总览、再看队列、最后看日志”的顺序进行排查。")
    add_bullets(doc, [
        "若评论列表为空，优先检查游戏分类、筛选条件和最近同步任务。",
        "若任务失败，进入任务队列查看错误信息与执行日志。",
        "若回复记录异常，检查审核队列、发送状态和原评论链接是否可访问。",
        "若导出失败，可先缩小筛选范围后再次尝试导出。",
    ])
    add_page_break(doc)

    add_heading(doc, "第16章  操作规范与最佳实践", 1)
    add_heading(doc, "16.1  推荐流程", 2)
    add_paragraph(doc, "为了提升团队协作效率，建议按固定顺序使用系统：先查看总览，再检查任务，再进入评论列表处理，最后回到回复记录页复核发送结果。")
    add_bullets(doc, [
        "工作开始前先查看总览页，了解当天评论压力和回复进展。",
        "定时查看任务队列，确保同步任务无持续异常。",
        "处理评论时优先聚焦自有产品与待处理分组。",
        "发送完成后及时复核回复记录，必要时登记删除需求或补充说明。",
    ])
    add_page_break(doc)

    add_heading(doc, "第17章  结语", 1)
    add_heading(doc, "17.1  使用说明", 2)
    add_paragraph(doc, f"本手册围绕{SOFTWARE_NAME}的核心模块，对登录、总览、游戏管理、评论处理、策略维护、同步任务、回复记录和用户管理等内容进行了逐页说明。")
    add_paragraph(doc, "实际使用过程中，可根据团队分工对页面使用频率进行调整，但建议始终保持数据分类、任务配置和角色权限的一致性，以保证系统运行稳定、操作边界清晰、记录可追溯。")


def main() -> None:
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    configure_styles(doc)
    reset_header_footer(doc)
    build_pages(doc)
    doc.save(str(OUTPUT))


if __name__ == "__main__":
    main()
